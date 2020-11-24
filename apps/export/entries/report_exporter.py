import os
import tempfile
import logging
from datetime import datetime
from subprocess import call

from django.conf import settings
from django.core.files.base import ContentFile, File
from django.db.models import (
    Case,
    When,
    Q,
    IntegerField
)
from docx.shared import Inches

from export.formats.docx import Document
from export.mime_types import (
    DOCX_MIME_TYPE,
    PDF_MIME_TYPE,
)

from analysis_framework.models import Widget
from entry.models import Entry, ExportData, Attribute, EntryGroupLabel
from entry.widgets import (
    scale_widget,
    time_widget,
    date_widget,
    time_range_widget,
    date_range_widget,
    geo_widget,
)
from lead.models import Lead
from utils.common import generate_filename
from tabular.viz import renderer as viz_renderer
from export.models import Export

logger = logging.getLogger(__name__)


class ExportDataVersionMismatch(Exception):
    pass


class ReportExporter:
    def __init__(self, exporting_widgets=None):
        self.doc = Document(
            os.path.join(settings.APPS_DIR, 'static/doc_export/template.docx')
        )
        self.lead_ids = []
        # ordered list of widget ids
        self.exporting_widgets = exporting_widgets or []
        self.region_data = {}
        # XXX: Limit memory usage? (Or use redis?)
        self.geoarea_data_cache = {}

    def load_exportables(self, exportables, regions):
        exportables = exportables.filter(
            data__report__levels__isnull=False,
        )

        self.exportables = exportables

        geo_data_required = Widget.objects.filter(
            id__in=self.exporting_widgets,
            widget_id=geo_widget.WIDGET_ID
        ).exists()
        # Load geo data if required
        if geo_data_required:
            self.region_data = {}
            for region in regions:
                # Collect geo area names for each admin level
                self.region_data[region.id] = [
                    {
                        'id': admin_level.id,
                        'level': admin_level.level,
                        'geo_area_titles': admin_level.get_geo_area_titles(),
                    }
                    for admin_level in region.adminlevel_set.all()
                ]

        return self

    def load_levels(self, levels):
        self.levels = levels
        return self

    def load_structure(self, structure):
        self.structure = structure
        return self

    def load_group_lables(self, entries, show_groups=False):
        self.entry_group_labels = {}
        if not show_groups:
            return self
        for entry, group, label in EntryGroupLabel.objects.filter(entry__in=entries).order_by().values_list(
                'entry_id', 'group__title', 'label__title'):
            entry_d = self.entry_group_labels[entry] = self.entry_group_labels.get(entry, [])
            entry_d.append([group, label])
        return self

    def load_text_from_text_widgets(self, entries, text_widget_ids):
        """
        Prefetch entry texts from text_Widgets
        text_widget_ids: [id1, id2, [id1, 'widget-key', 'widget-label']]
        """
        # User defined widget order (Generate order map)
        widget_map = {
            id: index
            for index, id in enumerate(text_widget_ids) if isinstance(id, int)
        }
        conditional_widget_map = {
            f'{array[0]}-{array[1]}': (index, array[2])
            for index, array in enumerate(text_widget_ids) if isinstance(array, list) and len(array) == 3
        }

        # Collect text from textWidgets for given entries using user defined widget data
        collected_widget_text = {}
        for attribute_d in Attribute.objects.filter(
            entry__in=entries,
            data__value__isnull=False,
            widget__id__in=[(id if isinstance(id, int) else id[0]) for id in text_widget_ids],
        ).values(
            'entry_id', 'data__value', 'widget__title', 'widget__id', 'widget__widget_id'
        ):
            entry_id = attribute_d['entry_id']
            widget_id = attribute_d['widget__id']
            widget_type = attribute_d['widget__widget_id']

            if widget_type == 'conditionalWidget':
                selected_widget_key = attribute_d['data__value'].get('selected_widget_key')
                widget_order, widget_title = (
                    conditional_widget_map.get(f'{widget_id}-{selected_widget_key}', (None, None))
                )
                text = (
                    (
                        attribute_d['data__value'].get(selected_widget_key) or {}
                    ).get('data') or {}
                ).get('value')
            else:
                widget_order = widget_map[widget_id]
                widget_title = attribute_d['widget__title']
                text = attribute_d['data__value']

            if widget_order is None:
                continue

            # Map text to entry->order->text
            collected_widget_text[entry_id] = collected_widget_text.get(entry_id, {})
            collected_widget_text[entry_id][widget_order] = [
                widget_title,
                text and text.strip(),  # Widget Text Value
            ]

        self.collected_widget_text = collected_widget_text
        return self

    def _generate_legend_page(self, project):
        para = self.doc.add_paragraph()
        self.legend_paragraph.add_next_paragraph(para)

        # todo in a table
        scale_widgets = project.analysis_framework.widget_set.filter(widget_id='scaleWidget')
        for widget in scale_widgets[::-1]:
            if not hasattr(widget, 'title'):
                continue
            title_para = self.doc.add_paragraph()
            title_para.ref.paragraph_format.right_indent = Inches(0.25)
            title_para.add_run(f'{widget.title}')
            for legend in widget.properties.get('data', {}).get('scale_units')[::-1]:
                para = self.doc.add_paragraph()
                para.ref.paragraph_format.right_indent = Inches(0.25)
                para.add_oval_shape(legend.get('color'))
                para.add_run(f'    {legend.get("label", "-Missing-")}')
                self.legend_paragraph.add_next_paragraph(para)
            self.legend_paragraph.add_next_paragraph(title_para)
        cond_widgets = project.analysis_framework.widget_set.filter(widget_id='conditionalWidget')
        for c_widget in cond_widgets[::-1]:
            for widget in filter(lambda x: x.get('widget', {}).get('widget_id') == 'scaleWidget', c_widget.properties.get('data', {}).get('widgets', [])):
                if not widget.get('widget', {}).get('title'):
                    continue
                title_para = self.doc.add_paragraph()
                title_para.ref.paragraph_format.right_indent = Inches(0.25)
                title_para.add_run(f'{widget.get("widget", {}).get("title")}')
                for legend in widget.get('widget', {}).get('properties', {}).get('data', {}).get('scale_units', [])[::-1]:
                    para = self.doc.add_paragraph()
                    para.ref.paragraph_format.right_indent = Inches(0.25)
                    para.add_oval_shape(legend.get('color'))
                    para.add_run(f'    {legend.get("label", "-Missing-")}')
                    self.legend_paragraph.add_next_paragraph(para)
                self.legend_paragraph.add_next_paragraph(title_para)

    def _add_scale_widget_data(self, para, data, bold=True):
        """
        report for scale widget expects following keys
            - title
            - label
            - color
        as described here: apps.entry.widgets.scale_widget._get_scale
        """
        if data.get('version') != scale_widget.DATA_VERSION:
            raise ExportDataVersionMismatch('Scale widget data is not upto date. '
                                            'Please wait, it will be updated soon.')
        if data.get('label', None) and data.get('color', None):
            para.add_run(', ', bold=True)
            para.add_oval_shape(data.get('color'))
            para.add_run('{}'.format(data.get('label', '')), bold)
            return True

    def _add_date_range_widget_data(self, para, data, bold=True):
        """
        report for date range widget expects following
            - tuple (from, to)
        as described here: apps.entry.widgets.date_range_widget._get_date
        """
        if data.get('version') != date_range_widget.DATA_VERSION:
            raise ExportDataVersionMismatch('Date Range widget data is not upto date. '
                                            'Please wait, it will be updated soon.')
        if len(data.get('values', [])) == 2 and any(data.get('values', [])):
            para.add_run(', {} - {}'.format(data['values'][0] or "00-00-00", data['values'][1] or "00-00-00"), bold)
            return True

    def _add_time_range_widget_data(self, para, data, bold=True):
        """
        report for time range widget expects following
            - tuple (from, to)
        as described here: apps.entry.widgets.time_range_widget._get_time
        """
        if data.get('version') != time_range_widget.DATA_VERSION:
            raise ExportDataVersionMismatch('Time Range widget data is not upto date. '
                                            'Please wait, it will be updated soon.')
        if len(data.get('values', [])) == 2 and any(data.get('values', [])):
            para.add_run(', {} - {}'.format(data['values'][0] or "~~:~~", data['values'][1] or "~~:~~"), bold)
            return True

    def _add_time_widget_data(self, para, data, bold=True):
        """
        report for time widget expects following
            - string (=time)
        as described here: apps.entry.widgets.time_widget._get_time
        """
        if data.get('version') != time_widget.DATA_VERSION:
            raise ExportDataVersionMismatch('Time widget data is not upto date. '
                                            'Please wait, it will be updated soon.')
        if data.get('value', None):
            para.add_run(', {}'.format(data['value']), bold)
            return True

    def _add_date_widget_data(self, para, data, bold=True):
        """
        report for date widget expects following
            - string (=date)
        as described here: apps.entry.widgets.date_widget
        """
        if data.get('version') != date_widget.DATA_VERSION:
            raise ExportDataVersionMismatch('Date widget data is not upto date. '
                                            'Please wait, it will be updated soon.')
        if data.get('value', None):
            para.add_run(', {}'.format(data['value']), bold)
            return True

    def _add_geo_widget_data(self, para, data, bold=True):
        # XXX: Cache this value.
        # Right now everything needs to be loaded so doing this at entry save can take lot of memory
        render_values = []
        geo_id_values = [str(v) for v in data.get('values') or []]

        if len(geo_id_values) == 0:
            return

        for region_id, admin_levels in self.region_data.items():
            for admin_level in admin_levels:
                geo_area_titles = admin_level['geo_area_titles']
                for geo_id in geo_id_values:
                    if geo_id not in geo_area_titles:
                        continue
                    if geo_id in self.geoarea_data_cache:
                        title = self.geoarea_data_cache[geo_id]
                        title and render_values.append(title)
                        continue
                    self.geoarea_data_cache[geo_id] = None

                    title = geo_area_titles[geo_id].get('title')
                    parent_id = geo_area_titles[geo_id].get('parent_id')
                    if admin_level['level'] == 1:
                        title and render_values.append(title)
                        self.geoarea_data_cache[geo_id] = title
                        continue

                    # Try to look through parent
                    for _level in range(0, admin_level['level'] - 1)[::-1]:
                        if parent_id:
                            _geo_area_titles = admin_levels[_level]['geo_area_titles']
                            _geo_area = _geo_area_titles.get(parent_id) or {}
                            _title = _geo_area.get('title')
                            parent_id = _geo_area.get('parent_id')
                            if _level == 1:
                                _title and render_values.append(_title)
                                self.geoarea_data_cache[geo_id] = title
                                break

        if render_values:
            para.add_run(f", {', '.join(set(render_values))}", bold)
            return True

    def _add_widget_information_into_report(self, para, report, bold=True):
        """
        based on widget annotate information into report

        :param para: Paragraph
        :param report: dict
        """
        if not isinstance(report, dict):
            return
        if 'widget_id' in report:
            widget_id = report.get('widget_id')
            mapper = {
                scale_widget.WIDGET_ID: self._add_scale_widget_data,
                date_range_widget.WIDGET_ID: self._add_date_range_widget_data,
                time_range_widget.WIDGET_ID: self._add_time_range_widget_data,
                time_widget.WIDGET_ID: self._add_time_widget_data,
                date_widget.WIDGET_ID: self._add_date_widget_data,
                geo_widget.WIDGET_ID: self._add_geo_widget_data,
            }
            if widget_id in mapper.keys():
                return mapper[widget_id](para, report, bold)

    def _generate_for_entry(self, entry):
        """
        Generate paragraphs for an entry
        """

        para = self.doc.add_paragraph().justify()

        para.add_run('[', bold=True)
        # Add lead-entry id
        url = '{protocol}://{site}/projects/{project}/leads/{lead}/entries/edit/?entry_id={entry}'.format(
            protocol=settings.HTTP_PROTOCOL,
            site=settings.DEEPER_FRONTEND_HOST,
            project=entry.lead.project.id,
            lead=entry.lead.id,
            entry=entry.id
        )
        # format of exporting_widgets = "[517,43,42,[405,"scalewidget-xe11vlcxs2gqdh1r","Scale"]]"
        widget_keys = [
            Widget.objects.get(id=each).key if not isinstance(each, list) else each[1]
            for each in self.exporting_widgets
        ]
        # para.add_hyperlink(url, f"{entry.lead.id}-{entry.id}")
        para.add_run(f'{entry.lead.id}-{entry.id}', bold=True)
        export_data = []
        for each in entry.exportdata_set.all():
            if 'other' in each.data.get('report', {}):
                for rep in each.data['report'].get('other', []):
                    if rep.get('widget_key') and rep['widget_key'] in widget_keys:
                        export_data.append(rep)
            else:
                export_datum = {**each.data.get('common', {}),
                                **each.data.get('report', {})}
                if export_datum.get('widget_key') and export_datum['widget_key'] in widget_keys:
                    export_data.append(export_datum)

        export_data.sort(key=lambda x: widget_keys.index(x['widget_key']))
        if export_data:
            for data in export_data:
                self._add_widget_information_into_report(para, data)
        para.add_run('] ', bold=True)

        # Format is
        # excerpt (source) OR excerpt \n text from widgets \n (source)
        # where source is hyperlinked to appropriate url

        # Excerpt can also be image
        excerpt = (
            entry.excerpt if entry.entry_type == Entry.EXCERPT
            else ''
        )
        para.add_run(excerpt)

        # Add texts from TextWidget
        widget_texts_exists = len(self.collected_widget_text.get(entry.id, [])) > 0
        entry_texts = self.collected_widget_text.get(entry.id, {})
        widget_texts_exists and self.doc.add_paragraph()  # Blank line
        for order in sorted(entry_texts.keys()):
            title, text = entry_texts[order]
            self.doc.add_heading(f'{title}: ', 1).add_run(text).bold = False
        if widget_texts_exists:
            para = self.doc.add_paragraph()

        # NOTE: Use doc.add_image for limiting image size to page width
        # and run.add_image for actual size image

        image = None
        image_text = None
        if entry.entry_type == Entry.IMAGE:
            image = entry.image
            # para.add_run().add_image(entry.image)
        elif entry.entry_type == Entry.DATA_SERIES and entry.tabular_field:
            image = viz_renderer.get_entry_image(entry)
            h_stats = (entry.tabular_field.cache or {}).get('health_stats', {})

            image_text = ' Total values: {}'.format(h_stats.get('total', 'N/A'))
            for key in ['invalid', 'null']:
                if h_stats.get(key):
                    image_text += f', {key.title()} values: {h_stats.get(key)}' if h_stats.get(key) else ''

        if image:
            self.doc.add_image(image)
            self.doc.add_paragraph(image_text).justify()
            para = self.doc.add_paragraph().justify()

        lead = entry.lead
        self.lead_ids.append(lead.id)

        source = lead.get_source_display() or 'Reference'
        author = lead.get_authors_display()
        url = lead.url or (
            lead.attachment and lead.attachment.get_file_url()
        )
        date = entry.lead.published_on

        para.add_run('(' if widget_texts_exists else ' (')

        # Add author is available
        (author and author.lower() != (source or '').lower()) and (
            para.add_hyperlink(url, f'{author}, ') if url else para.add_run(f'{author}, ')
        )
        # Add source (with url if available)
        para.add_hyperlink(url, source) if url else para.add_run(source)
        # Add (confidential) to source without ,
        lead.confidentiality == Lead.CONFIDENTIAL and para.add_run(' (confidential)')
        # Add lead title if available
        lead.title and para.add_run(f", {lead.title}")
        # Finally add date
        # TODO: use utils.common.format_date and perhaps use information date
        date and para.add_run(f", {date.strftime('%d/%m/%Y')}")
        # para.add_run(f", {'Verified' if entry.verified else 'Unverified'}")
        para.add_run(')')
        # para = self.doc.add_paragraph().justify()

        # Adding Entry Group Labels
        group_labels = self.entry_group_labels.get(entry.pk) or []
        if len(group_labels) > 0:
            para.add_run(' (')
            para.add_run(', '.join([f'{group} : {label}' for group, label in group_labels]))
            para.add_run(')')

        self.doc.add_paragraph()

    def _load_into_levels(
            self,
            entry,
            keys,
            levels,
            entries_map,
            valid_levels,
    ):
        """
        Map an entry into corresponding levels
        """
        parent_level_valid = False
        for level in levels:
            level_id = level.get('id')
            valid_level = (level_id in keys)

            if valid_level:
                if not entries_map.get(level_id):
                    entries_map[level_id] = []
                entries_map[level_id].append(entry)

            sublevels = level.get('sublevels')
            if sublevels:
                valid_level = valid_level or self._load_into_levels(
                    entry,
                    keys,
                    sublevels,
                    entries_map,
                    valid_levels,
                )

            if valid_level and level_id not in valid_levels:
                valid_levels.append(level_id)

            parent_level_valid = parent_level_valid or valid_level
        return parent_level_valid

    def _generate_for_levels(
            self,
            levels,
            level_entries_map,
            valid_levels,
            structures=None,
            heading_level=2,
    ):
        """
        Generate paragraphs for all entries in this level and recursively
        do it for further sublevels
        """
        if structures is not None:
            level_map = dict((level.get('id'), level) for level in levels)
            levels = [level_map[s['id']] for s in structures]

        for level in levels:
            if level.get('id') not in valid_levels:
                continue

            title = level.get('title')
            entries = level_entries_map.get(level.get('id'))
            sublevels = level.get('sublevels')

            if entries or sublevels:
                self.doc.add_heading(title, heading_level)
                self.doc.add_paragraph()

            if entries:
                [self._generate_for_entry(entry) for entry in entries]

            if sublevels:
                substructures = None
                if structures:
                    substructures = next((
                        s.get('levels') for s in structures
                        if s['id'] == level.get('id')
                    ), None)

                self._generate_for_levels(
                    sublevels,
                    level_entries_map,
                    valid_levels,
                    substructures,
                    heading_level + 1,
                )

    def _generate_for_uncategorized(self, entries):
        entries = entries.exclude(
            Q(exportdata__data__report__keys__isnull=False) |
            Q(exportdata__data__report__keys__len__gt=0)
        )
        if entries.count() == 0:
            return

        self.doc.add_heading('Uncategorized', 2)
        self.doc.add_paragraph()

        for entry in entries:
            self._generate_for_entry(entry)

    def pre_build_document(self, project):
        """
        Structure the document
        """
        self.doc.add_heading('DEEP Export — {} — {}'.format(datetime.today().strftime('%b %d, %Y'),
                                                            project.title),
                             1)
        self.doc.add_paragraph()

        self.legend_heading = self.doc.add_heading('Legends', 2)
        self.legend_paragraph = self.doc.add_paragraph()

    def add_entries(self, entries):
        """
        Add entries and generate parapgraphs for all entries
        """
        if entries:
            self.pre_build_document(entries[0].project)
        exportables = self.exportables
        af_levels_map = dict((str(level.get('id')), level.get('levels')) for level in self.levels)
        uncategorized = False

        if self.structure:
            ids = [s['id'] for s in self.structure]
            uncategorized = 'uncategorized' in ids
            ids = [id for id in ids if id != 'uncategorized']

            order = Case(*[
                When(pk=pk, then=pos)
                for pos, pk
                in enumerate(ids)
            ])
            exportables = exportables.filter(pk__in=ids).order_by(order)

        for exportable in exportables:
            levels = (
                # Custom levels provided by client
                af_levels_map.get(str(exportable.pk)) or
                # Predefined levels available in server
                exportable.data.get('report').get('levels')
            )

            level_entries_map = {}
            valid_levels = []
            for entry in entries:
                # TODO
                # Set entry.report_data to all exportdata for all exportable
                # for this entry for later use

                export_data = ExportData.objects.filter(
                    entry=entry,
                    exportable=exportable,
                    data__report__keys__isnull=False,
                ).first()

                if export_data:
                    self._load_into_levels(
                        entry, export_data.data.get('report').get('keys'),
                        levels, level_entries_map, valid_levels,
                    )

            structures = self.structure and next((
                s.get('levels') for s in self.structure
                if str(s['id']) == str(exportable.id)
            ), None)
            self._generate_for_levels(levels, level_entries_map,
                                      valid_levels, structures)

        if uncategorized:
            self._generate_for_uncategorized(entries)
        if entries:
            self._generate_legend_page(entries[0].project)

        return self

    def export(self, pdf=False):
        """
        Export and return export data
        """

        # Get all leads to generate Bibliography
        lead_ids = list(set(self.lead_ids))
        leads = Lead.objects.filter(id__in=lead_ids)

        self.doc.add_paragraph().add_horizontal_line()
        self.doc.add_paragraph()
        self.doc.add_heading('Bibliography', 1)
        self.doc.add_paragraph()

        for lead in leads:
            # Format is"
            # Source. Title. Date. Url

            para = self.doc.add_paragraph()
            author = lead.get_authors_display()
            source = lead.get_source_display() or 'Missing source'

            author and para.add_run(f'{author}.')
            para.add_run(f' {source}.')
            para.add_run(f' {lead.title}.')
            lead.published_on and para.add_run(f" {lead.published_on.strftime('%m/%d/%y')}. ")

            para = self.doc.add_paragraph()
            url = lead.url or (
                lead.attachment and lead.attachment.get_file_url()
            )
            if url:
                para.add_hyperlink(url, url)
            else:
                para.add_run('Missing url.')

            if lead.confidentiality == Lead.CONFIDENTIAL:
                para.add_run(' (confidential)')

            self.doc.add_paragraph()
        # self.doc.add_page_break()

        if pdf:
            temp_doc = tempfile.NamedTemporaryFile(dir=settings.TEMP_DIR)
            self.doc.save_to_file(temp_doc)

            filename = temp_doc.name.split('/')[-1]
            temp_pdf = os.path.join(settings.TEMP_DIR,
                                    '{}.pdf'.format(filename))

            call(['libreoffice', '--headless', '--convert-to',
                  'pdf', temp_doc.name, '--outdir', settings.TEMP_DIR])

            filename = generate_filename('Entries General Export', 'pdf')
            file = File(open(temp_pdf, 'rb'))
            export_format = Export.PDF
            export_mime_type = PDF_MIME_TYPE

            # Cleanup
            os.unlink(temp_pdf)
            temp_doc.close()

        else:
            buffer = self.doc.save()

            filename = generate_filename('Entries General Export', 'docx')
            file = ContentFile(buffer)
            export_format = Export.DOCX
            export_mime_type = DOCX_MIME_TYPE

        return filename, export_format, export_mime_type, file
